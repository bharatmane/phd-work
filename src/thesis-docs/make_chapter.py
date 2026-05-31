"""
Converts a thesis chapter markdown file to a Word docx using the Alliance University template.
Usage: python make_chapter.py <input.md> <output.docx>
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

TEMPLATE = r'C:\Users\bhara\template.docx'

# ── Style mapping ─────────────────────────────────────────────────────────────
# MD level  →  template style name
HEADING_STYLES = {
    1: "Heading 1",   # ## 1.1
    2: "Heading 2",   # ### 1.1.1
    3: "Heading 3",   # ####
}
BODY_STYLE      = "Chapter Body Text"
LIST_STYLE      = "List Paragraph"
CODE_FONT       = "Courier New"
CODE_SIZE       = Pt(10)

# ── Inline parser ─────────────────────────────────────────────────────────────

def parse_inline(text):
    """
    Returns list of (segment_text, bold, italic, code) tuples.
    Handles **bold**, *italic*, `code`, and combinations.
    """
    segments = []
    # Pattern: **bold**, *italic*, `code`  (order matters — ** before *)
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)')
    last = 0
    for m in pattern.finditer(text):
        start, end = m.span()
        if start > last:
            segments.append((text[last:start], False, False, False))
        if m.group(1).startswith('**'):
            segments.append((m.group(2), True, False, False))
        elif m.group(1).startswith('*'):
            segments.append((m.group(3), False, True, False))
        else:
            segments.append((m.group(4), False, False, True))
        last = end
    if last < len(text):
        segments.append((text[last:], False, False, False))
    return segments if segments else [(text, False, False, False)]


def add_formatted_para(doc, text, style_name):
    """Add a paragraph with inline formatting applied."""
    available = [s.name for s in doc.styles]
    sty = style_name if style_name in available else "Normal"
    para = doc.add_paragraph(style=sty)
    segments = parse_inline(text)
    for seg_text, bold, italic, code in segments:
        run = para.add_run(seg_text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if code:
            run.font.name = CODE_FONT
            run.font.size = CODE_SIZE
    return para


# ── Main conversion ───────────────────────────────────────────────────────────

def convert(md_path, out_path):
    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document(TEMPLATE)

    # Remove all existing content from template
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(doc.tables):
        t._element.getparent().remove(t._element)

    available_styles = [s.name for s in doc.styles]

    def sty(name):
        return name if name in available_styles else "Normal"

    in_code_block = False
    code_lines    = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()

        # ── Code block ────────────────────────────────────────────────────────
        if stripped.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                # End of code block — emit as body text with Courier
                if code_lines:
                    p = doc.add_paragraph(style=sty(BODY_STYLE))
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = CODE_FONT
                    run.font.size = CODE_SIZE
                in_code_block = False
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Skip dividers and end-of-chapter notes ────────────────────────────
        if stripped == '---' or stripped.startswith('*End of') or stripped.startswith('*Chapter'):
            i += 1
            continue

        # ── Empty line ────────────────────────────────────────────────────────
        if stripped == '':
            i += 1
            continue

        # ── Chapter heading  # CHAPTER N: TITLE ──────────────────────────────
        if line.startswith('# '):
            heading = line[2:].strip()
            if ':' in heading:
                ch_num, ch_title = heading.split(':', 1)
            else:
                ch_num, ch_title = heading, ''
            # Chapter Number paragraph
            p = doc.add_paragraph(style=sty("Chapter Number"))
            p.add_run(ch_num.strip())
            # Chapter Title paragraph
            if ch_title.strip():
                p = doc.add_paragraph(style=sty("Chapter Title"))
                p.add_run(ch_title.strip())
            i += 1
            continue

        # ── Section headings ──────────────────────────────────────────────────
        if line.startswith('#### '):
            add_formatted_para(doc, line[5:].strip(), sty(HEADING_STYLES[3]))
            i += 1
            continue

        if line.startswith('### '):
            add_formatted_para(doc, line[4:].strip(), sty(HEADING_STYLES[2]))
            i += 1
            continue

        if line.startswith('## '):
            add_formatted_para(doc, line[3:].strip(), sty(HEADING_STYLES[1]))
            i += 1
            continue

        # ── Bullet list ───────────────────────────────────────────────────────
        if stripped.startswith('- ') or stripped.startswith('* '):
            item_text = stripped[2:].strip()
            add_formatted_para(doc, item_text, sty(LIST_STYLE))
            i += 1
            continue

        # ── Body text ─────────────────────────────────────────────────────────
        add_formatted_para(doc, stripped, sty(BODY_STYLE))
        i += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")
    # Summary
    headings = [p.text[:70] for p in doc.paragraphs
                if p.style.name in ("Chapter Number","Chapter Title","Heading 1","Heading 2","Heading 3")]
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print("Structure:")
    for h in headings:
        print(f"  {h}")


# ── Entry point ───────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python make_chapter.py <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
