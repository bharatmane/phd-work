"""
Creates 02_prelim_pages.docx by filling in the Alliance University thesis template
with Bharat's actual content. Uses the template directly so all styles are native.
"""
import shutil, copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

TEMPLATE = r'C:\Users\bhara\template.docx'
OUTPUT   = r'C:\Users\bhara\OneDrive\Documents\Claude\Projects\Phd\Thesis\DocxFiles\02_prelim_pages.docx'

THESIS_TITLE = (
    "Explainable Deep Learning for Multi-Level Program Comprehension: "
    "Identifier Readability, Code Snippet Analysis, and Developer Experience Classification"
)

# ─────────────────────────────────────────────────────────────────────────────
# Content blocks
# ─────────────────────────────────────────────────────────────────────────────

DECLARATION_TEXT = (
    "I declare that the thesis entitled "
    f"“{THESIS_TITLE}” "
    "has been prepared by me under the guidance of Dr. Rathnakar Achary, "
    "Associate Professor, Alliance School of Advance Computing, Alliance University. "
    "No part of this thesis has formed the basis for the award of any degree in any "
    "university or fellowship previously. The work reported in this thesis was carried "
    "out independently by me unless explicitly acknowledged in the text."
)

CERTIFICATE_TEXT = (
    "I certify that Bharat Babaso Mane has prepared his thesis entitled "
    f"“{THESIS_TITLE}”, "
    "for the award of the PhD degree of Alliance University, under my guidance. "
    "He has carried out the work at the Alliance School of Advance Computing, "
    "Alliance University."
)

DEDICATION_PARAS = [
    "To my parents, whose sacrifices made every page of this work possible.",
    "To my supervisor, Dr. Rathnakar Achary, for the direction, patience, and rigour that shaped this research.",
    "And to every software developer who has ever inherited someone else’s code and wondered what on earth it was supposed to do.",
]

ACKNOWLEDGEMENT_PARAS = [
    "Dr. Rathnakar Achary, my supervisor, offered something rarer than technical guidance: he consistently asked better questions than I had thought to ask myself. His insistence that results must be explainable — not just accurate — gave this thesis its defining theme.",
    "The faculty and staff of the Alliance School of Advance Computing provided an environment in which research could be taken seriously. I am grateful for access to computational resources, library systems, and collegial exchange that sustained this work.",
    "My colleagues and peers across three years of doctoral study offered critical feedback, pointed out my blind spots, and were generous enough to disagree with me when they had reason to.",
    "My family understood that a PhD is not a project with weekends. Their patience with my absences — physical and mental — is something I recognise and can only partially repay.",
    "The datasets that made this research possible were created and shared by Paakhim10 (Kaggle Code Snippets: Insights and Readability) and by Perez, Urtado, and Vauttier (Zenodo developer experience dataset). Open data is a public good, and this thesis would not exist without it.",
    "Bharat Babaso Mane\nBengaluru, 2026",
]

ABSTRACT_PARAS = [
    "Program comprehension — the cognitive process by which software developers understand what code does, how it is structured, and who built it — underpins nearly every phase of the software development lifecycle. Maintenance, debugging, refactoring, code review, and developer onboarding all depend on a developer’s ability to read and make sense of source code quickly and accurately. Industry estimates consistently place maintenance at sixty to seventy percent of total software lifecycle cost, and a substantial portion of that cost is attributable to code that is difficult to read.",
    "The first study, IRAF-XADL (Identifier Readability Analysis Framework using Explainable Attention-Based Deep Learning), operates at the finest grain: the individual identifier. IRAF-XADL extracts identifiers using language-specific abstract syntax tree parsers for Python (LibCST) and C++ (Tree-Sitter), computes ten linguistically and cognitively grounded readability parameters from each identifier, encodes it with CodeBERT contextual embeddings, and classifies it as High, Medium, or Low readability using a Self-Attention BiLSTM optimised with AdamW. SHAP-based post-hoc explanations identify which features drove each prediction. On the Code Snippets: Insights and Readability benchmark, IRAF-XADL achieves test accuracy of 97.36% for Python and 97.94% for C++, surpassing the best published baseline by more than fifteen percentage points.",
    "The second study, ECRVR-MVEL (Explainable Code Readability Classification Using Vector Representations and Majority Voting-Based Ensemble Learning), widens the unit of analysis from a single identifier to a complete code snippet. ECRVR-MVEL applies CodeBERT to encode the full snippet, then combines three structurally diverse classifiers — a Graph Convolutional Network, a Deep Belief Network, and a Bidirectional Temporal Convolutional Network — in a weighted majority vote, optimised with the Nadam algorithm. LIME provides local explanations for each prediction. On the same benchmark dataset, ECRVR-MVEL achieves test accuracy of 98.15% for Python and 98.38% for C++, outperforming the Neural Network baseline by more than eight percentage points.",
    "The third study, EESQA-DELMOA (Empirical Evaluation of Software Quality Assessment through Developer Experience Level Using Metaheuristic Optimisation Algorithms), shifts focus from the code artefact to the developer who produced it. EESQA-DELMOA applies min-max normalisation, Bio-inspired Artificial Hummingbird Behaviour (BAHB) feature selection, a Simplified Spiking Neural Network (SSNN) classifier, and Adaptive Migration Butterfly Optimisation Algorithm (AMBOA) hyperparameter tuning to classify developer experience into six categories. On the Perez et al. dataset (703 developer profiles), the system achieves test accuracy of 98.74% with an execution time of 8.27 seconds — the lowest of all compared methods.",
    "A cross-study analysis reveals a consistent finding: SHAP explanations in Study 1 and LIME explanations in Study 2 independently identify Meaningful Clarity (MC) and Naming Conformance (NC) as the dominant predictors of code readability. This convergence, across two independent explainability methods and two levels of analysis, validates the relevance of the ten-parameter feature set beyond the identifier level.",
    "Together, the three studies establish that program comprehension can be assessed automatically, accurately, and interpretably at every level of the software development artefact hierarchy — from the naming of a single variable, through the structure of a code snippet, to the experience of the developer who wrote it.",
    "Keywords: Program Comprehension; Code Readability; Identifier Quality; CodeBERT; Self-Attention BiLSTM; Graph Convolutional Network; Spiking Neural Network; Explainable AI; SHAP; LIME; Software Quality Assessment; Developer Experience",
]

PREFACE_PARAS = [
    "This thesis reports three independent but related studies completed as part of a doctoral research programme in computer science and software engineering at Alliance University, Bengaluru. Each study was designed to address a specific gap in the automated assessment of program comprehension and source code quality.",
    "The studies are presented in order of increasing abstraction: from the individual identifier (Chapter 3) to the code snippet (Chapter 4) to the developer behind the code (Chapter 5). Chapters 3, 4, and 5 are substantially based on papers that have been peer-reviewed and accepted for publication. The bibliographic details of these publications appear in the List of Publications at the end of the thesis. Wherever results, figures, or tables are drawn directly from the published papers, this is indicated in the text. The comparative analysis, cross-study discussion, and synthesis in Chapters 6 and 7 are original to this thesis and do not appear in the published papers.",
]

TABLES_ENTRIES = [
    "Table 3.1\tTen readability parameters: name, description, and cognitive basis\t3",
    "Table 3.2\tHyperparameter configuration of the SA-BiLSTM model\t3",
    "Table 3.3\tDataset statistics: Python and C++ code snippets\t3",
    "Table 3.4\tIRAF-XADL results on Python data (70/30 split)\t3",
    "Table 3.5\tIRAF-XADL results on C++ data (70/30 split)\t3",
    "Table 3.6\tComparative analysis: IRAF-XADL vs. baselines\t3",
    "Table 3.7\tFeature pairwise correlation matrix (Pearson r)\t3",
    "Table 3.8\tAblation results on Python test set\t3",
    "Table 3.9\tConfusion matrix — Python test set\t3",
    "Table 4.1\tNadam hyperparameter configuration\t4",
    "Table 4.2\tDataset statistics (snippet level)\t4",
    "Table 4.3\tIndividual and ensemble results — Python, 70% training\t4",
    "Table 4.4\tIndividual and ensemble results — Python, 30% testing\t4",
    "Table 4.5\tPer-class breakdown — Python, 30% testing\t4",
    "Table 4.6\tIndividual and ensemble results — C++, 70% training\t4",
    "Table 4.7\tIndividual and ensemble results — C++, 30% testing\t4",
    "Table 4.8\tECRVR-MVEL vs. baselines — Python\t4",
    "Table 4.9\tECRVR-MVEL vs. baselines — C++\t4",
    "Table 4.10\tAblation conditions for ECRVR-MVEL\t4",
    "Table 4.11\tAblation results on Python test set\t4",
    "Table 4.12\tPairwise diversity on Python test set\t4",
    "Table 4.13\tLIME stability across 10 runs\t4",
    "Table 5.1\tDataset statistics — developer experience\t5",
    "Table 5.2\tEESQA-DELMOA classification results (70/30 split)\t5",
    "Table 5.3\tEESQA-DELMOA vs. baseline classifiers\t5",
    "Table 5.4\tExecution time comparison\t5",
    "Table 5.5\tStage-by-stage execution time breakdown\t5",
    "Table 5.6\tAblation conditions and results\t5",
    "Table 5.7\tFeature selection frequency across 20 BAHB runs\t5",
    "Table 5.8\tPer-class results on test set\t5",
    "Table 5.9\tImpact of imbalance mitigation strategies\t5",
    "Table 6.1\tPerformance summary across three studies\t6",
    "Table 6.2\tEnsemble vs. individual classifier performance\t6",
]

FIGURES_ENTRIES = [
    "Figure 1.1\tThe three-level program comprehension hierarchy addressed in this thesis\t1",
    "Figure 1.2\tThesis organisation and inter-chapter relationships\t1",
    "Figure 3.1\tOverall architecture of the IRAF-XADL framework\t3",
    "Figure 3.2\tArchitecture of the self-attention mechanism\t3",
    "Figure 3.3\tSA-BiLSTM model diagram\t3",
    "Figure 3.4\tConfusion matrices and PR/ROC curves: IRAF-XADL on Python\t3",
    "Figure 3.5\tTraining and validation accuracy curves: IRAF-XADL on Python\t3",
    "Figure 3.6\tTraining and validation loss curves: IRAF-XADL on Python\t3",
    "Figure 3.7\tConfusion matrices and PR/ROC curves: IRAF-XADL on C++\t3",
    "Figure 3.8\tTraining and validation accuracy curves: IRAF-XADL on C++\t3",
    "Figure 3.9\tSHAP feature importance: Python data\t3",
    "Figure 3.10\tSHAP feature importance: C++ data\t3",
    "Figure 4.1\tOverall architecture of the ECRVR-MVEL framework\t4",
    "Figure 4.2\tGCN, DBN, Bi-TCN confusion matrices and ROC curves: Python\t4",
    "Figure 4.3\tEnsemble accuracy curve: Python\t4",
    "Figure 4.4\tEnsemble loss curve: Python\t4",
    "Figure 4.5\tGCN, DBN, Bi-TCN confusion matrices and ROC curves: C++\t4",
    "Figure 4.6\tEnsemble accuracy curve: C++\t4",
    "Figure 4.7\tLIME explanations: Python (Low, Medium, High)\t4",
    "Figure 4.8\tLIME explanations: C++ (Low, Medium, High)\t4",
    "Figure 5.1\tOverall architecture of the EESQA-DELMOA framework\t5",
    "Figure 5.2\tSSNN structure\t5",
    "Figure 5.3\tAccuracy curve: EESQA-DELMOA\t5",
    "Figure 5.4\tLoss curve: EESQA-DELMOA\t5",
    "Figure 5.5\tPR curve: EESQA-DELMOA\t5",
    "Figure 5.6\tROC curve: EESQA-DELMOA\t5",
    "Figure 5.7\tComparative accuracy: EESQA-DELMOA vs. baselines\t5",
    "Figure 5.8\tExecution time comparison\t5",
    "Figure 6.1\tCross-study relationship diagram\t6",
    "Figure 6.2\tFeature importance convergence: SHAP (Study 1) and LIME (Study 2)\t6",
]

APPENDIX_ENTRIES = [
    ("APPENDIX A", "Dataset Details — Code Snippets: Insights and Readability (Kaggle) and Developer Experience Dataset (Zenodo)"),
    ("APPENDIX B", "Hyperparameter Tables — IRAF-XADL, ECRVR-MVEL, and EESQA-DELMOA"),
    ("APPENDIX C", "Evaluation Metric Formulas"),
    ("APPENDIX D", "Reproducibility Information — Library Versions, Random Seeds, and Hardware"),
]

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def clear_and_set(para, text):
    """Replace all runs in a paragraph with a single run containing text."""
    # Save formatting from first run if available
    bold = None
    size = None
    fname = None
    if para.runs:
        r0 = para.runs[0]
        bold  = r0.bold
        size  = r0.font.size
        fname = r0.font.name
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
        if bold  is not None: para.runs[0].bold = bold
        if size  is not None: para.runs[0].font.size = size
        if fname is not None: para.runs[0].font.name = fname
    else:
        para.add_run(text)

def insert_para_after(ref_para, text, style_name):
    """Insert a new paragraph (same style) with text immediately after ref_para."""
    new_p = copy.deepcopy(ref_para._p)
    ref_para._p.addnext(new_p)
    # Find the new paragraph object
    parent = ref_para._p.getparent()
    idx = list(parent).index(new_p)
    # Clear runs in new element
    from docx.oxml.ns import qn
    for r in new_p.findall(qn('w:r')):
        new_p.remove(r)
    # We'll access it as a paragraph object — re-parse
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, ref_para._p.getparent())
    run = new_para.add_run(text)
    # inherit font from template para
    if ref_para.runs:
        r0 = ref_para.runs[0]
        if r0.font.name: run.font.name = r0.font.name
        if r0.font.size: run.font.size = r0.font.size
    return new_para

def delete_para(para):
    p = para._p
    p.getparent().remove(p)

# ─────────────────────────────────────────────────────────────────────────────
# Open template and work on it
# ─────────────────────────────────────────────────────────────────────────────

doc = Document(TEMPLATE)
paras = doc.paragraphs   # live list

# ── Step 1: Remove title page section (paras 0-20) ───────────────────────────
# We delete from the END to avoid index shifting
# Title page runs from index 0 to 20 inclusive (21 paragraphs)
# The section break (para 21) is a blank Head Introductory Pages — we keep it
# as the first paragraph to ensure proper section formatting.

# Delete paras 0-20 (reverse order)
to_delete = list(doc.paragraphs[:21])
for p in to_delete:
    delete_para(p)

# Refresh paragraph list
paras = doc.paragraphs

# ── Step 2: Fill Declaration (paras 0=blank, 1=DECLARATION, 2=text, ...) ─────
# After deleting title page, paras are renumbered.
# Let's work by finding paragraphs by their current text content.

def find_para(text_contains):
    for p in doc.paragraphs:
        if text_contains.lower() in p.text.lower():
            return p
    return None

# DECLARATION body text
p = find_para("I declare that the thesis")
if p:
    clear_and_set(p, DECLARATION_TEXT)

# Scholar name under declaration
p = find_para("Scholar")
if p and "Name" in p.text:
    clear_and_set(p, "Bharat Babaso Mane")

# School under declaration
for para in doc.paragraphs:
    if "School / College Name" in para.text and "Alliance University" in para.text:
        clear_and_set(para, "Alliance School of Advance Computing, Alliance University")
        break   # only first occurrence (declaration)

# ── Step 3: Fill Certificate ─────────────────────────────────────────────────

p = find_para("I certify that")
if p:
    clear_and_set(p, CERTIFICATE_TEXT)

# Guide name — find [Guide's Name] or [name of supervisor] placeholder
for para in doc.paragraphs:
    if "Guide" in para.text and ("name" in para.text.lower() or "�" in para.text):
        clear_and_set(para, "Dr. Rathnakar Achary")
        break

# All remaining [School / College Name] occurrences → guide's school
for para in doc.paragraphs:
    if "School / College Name" in para.text:
        clear_and_set(para, "Alliance School of Advance Computing, Alliance University")

# ── Step 4: Fill Dedication ───────────────────────────────────────────────────

p = find_para("Dedication body text")
if p:
    clear_and_set(p, DEDICATION_PARAS[0])
    prev = p
    for line in DEDICATION_PARAS[1:]:
        prev = insert_para_after(prev, line, p.style.name)

# ── Step 5: Fill Acknowledgement ─────────────────────────────────────────────

p = find_para("I Acknowledge")
if p:
    clear_and_set(p, ACKNOWLEDGEMENT_PARAS[0])
    prev = p
    for line in ACKNOWLEDGEMENT_PARAS[1:]:
        prev = insert_para_after(prev, line, p.style.name)

# ── Step 6: Fill Abstract ─────────────────────────────────────────────────────

# There are TWO paragraphs with "[Abstract text starts here]" — first is Abstract,
# second is Preface. Handle them in order.
abstract_done = False
preface_done  = False

for para in doc.paragraphs:
    if "Abstract text starts here" in para.text or "abstract text starts here" in para.text.lower():
        if not abstract_done:
            clear_and_set(para, ABSTRACT_PARAS[0])
            prev = para
            for line in ABSTRACT_PARAS[1:]:
                prev = insert_para_after(prev, line, para.style.name)
            abstract_done = True
        elif not preface_done:
            clear_and_set(para, PREFACE_PARAS[0])
            prev = para
            for line in PREFACE_PARAS[1:]:
                prev = insert_para_after(prev, line, para.style.name)
            preface_done = True

# ── Step 7: Fill List of Tables ───────────────────────────────────────────────

p = find_para("[List starts here]")
if p:
    clear_and_set(p, TABLES_ENTRIES[0])
    prev = p
    for entry in TABLES_ENTRIES[1:]:
        prev = insert_para_after(prev, entry, p.style.name)

# ── Step 8: Fill List of Figures ─────────────────────────────────────────────
# After List of Tables was filled, the placeholder is gone.
# Find the remaining "[List starts here]" (which is now under LIST OF FIGURES).
for para in doc.paragraphs:
    txt = para.text.strip()
    if "List starts here" in txt:
        clear_and_set(para, FIGURES_ENTRIES[0])
        prev = para
        for entry in FIGURES_ENTRIES[1:]:
            prev = insert_para_after(prev, entry, para.style.name)
        break

# ── Step 9: Fill List of Appendices ──────────────────────────────────────────

for para in doc.paragraphs:
    if "APPENDIX A" in para.text and "Type your" in para.text:
        label, desc = APPENDIX_ENTRIES[0]
        clear_and_set(para, f"{label}:\t{desc}")
        break

for para in doc.paragraphs:
    if "APPENDIX B" in para.text and "Type your" in para.text:
        label, desc = APPENDIX_ENTRIES[1]
        clear_and_set(para, f"{label}:\t{desc}")
        break

for para in doc.paragraphs:
    if "APPENDIX C" in para.text and "Type your" in para.text:
        label, desc = APPENDIX_ENTRIES[2]
        clear_and_set(para, f"{label}:\t{desc}")
        break

for para in doc.paragraphs:
    if "APPENDIX D" in para.text and "Type your" in para.text:
        label, desc = APPENDIX_ENTRIES[3]
        clear_and_set(para, f"{label}:\t{desc}")
        break

# ── Step 10: Remove everything after the last Appendix entry ─────────────────
# The prelim section ends after the APPENDIX D line.
# Everything from "CHAPTER 1" / "Chapter Number" style onward is chapters.

# Find the last paragraph that is clearly a prelim-page item.
# Strategy: find index of last APPENDIX line, then delete everything after it.

last_appendix_idx = None
for i, para in enumerate(doc.paragraphs):
    if "APPENDIX" in para.text and "Type your" not in para.text:
        last_appendix_idx = i

if last_appendix_idx is not None:
    to_delete = list(doc.paragraphs[last_appendix_idx + 1:])
    for p in to_delete:
        try:
            delete_para(p)
        except Exception:
            pass

# ── Save ──────────────────────────────────────────────────────────────────────

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print("\nSection headings found:")
for p in doc.paragraphs:
    if p.style.name in ("Head Introductory Pages", "Chapter Title"):
        print(f"  [{p.style.name}] {p.text}")
