"""
Creates 01_Title.docx using the Alliance University thesis template styles.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

TEMPLATE = r'C:\Users\bhara\template.docx'
OUTPUT   = r'C:\Users\bhara\OneDrive\Documents\Claude\Projects\Phd\Thesis\DocxFiles\01_Title.docx'

THESIS_TITLE = (
    "EXPLAINABLE DEEP LEARNING FOR MULTI-LEVEL PROGRAM COMPREHENSION: "
    "IDENTIFIER READABILITY, CODE SNIPPET ANALYSIS, AND "
    "DEVELOPER EXPERIENCE CLASSIFICATION"
)

# ── helpers ──────────────────────────────────────────────────────────────────

def set_para_text(para, text, bold=None, font_size=None, font_name=None):
    """Clear runs and set a single run with the given text and formatting."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        run = para.runs[0]
    else:
        run = para.add_run()
    run.text = text
    if bold is not None:
        run.bold = bold
    if font_size is not None:
        run.font.size = Pt(font_size)
    if font_name is not None:
        run.font.name = font_name

def add_para(doc, text, style_name, bold=None, size=None, font="Times New Roman",
             align=WD_ALIGN_PARAGRAPH.CENTER, space_before=None, space_after=None):
    """Add a paragraph with the given style if it exists, else use Normal."""
    style_names = [s.name for s in doc.styles]
    sty = style_name if style_name in style_names else "Normal"
    p = doc.add_paragraph(style=sty)
    p.alignment = align
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = font
    if bold is not None:
        run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    return p

# ── open template (to inherit styles and page setup) ─────────────────────────

doc = Document(TEMPLATE)

# Remove ALL existing paragraphs (we'll rebuild the title page from scratch)
# Keep only the first section's page setup
for para in list(doc.paragraphs):
    p_elem = para._element
    p_elem.getparent().remove(p_elem)

# Also clear tables if any leaked in
for tbl in list(doc.tables):
    tbl._element.getparent().remove(tbl._element)

# ── build title page ──────────────────────────────────────────────────────────

# Get available style names for safe lookup
style_names = [s.name for s in doc.styles]

def sty(name):
    return name if name in style_names else "Normal"

# ── 1. Thesis Title ───────────────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Title"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(72)   # push down from top
p.paragraph_format.space_after  = Pt(24)
run = p.add_run(THESIS_TITLE)
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(14)

# ── 2. "Thesis submitted to Alliance University" ──────────────────────────────
p = doc.add_paragraph(style=sty("Text Thesis submitted to"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run("Thesis submitted to Alliance University")
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# ── 3. "for the Award of" ──────────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text Thesis 2"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(12)
run = p.add_run("for the Award of")
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# ── 4. "DOCTOR OF PHILOSOPHY" ────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text 2 Doctor of Philosophy"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
run = p.add_run("DOCTOR OF PHILOSOPHY")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(14)

# ── 5. "In" ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text Thesis 2"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run("In")
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# ── 6. Faculty ────────────────────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text 2 Doctor of Philosophy"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(18)
run = p.add_run("FACULTY OF ENGINEERING AND TECHNOLOGY")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(13)

# ── 7. "By" ───────────────────────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text Thesis 2"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run("By")
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# ── 8. Scholar name ───────────────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text 2 Doctor of Philosophy"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run("BHARAT BABASO MANE")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(13)

# ── 9. Registration number ────────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text Thesis 2"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(18)
run = p.add_run("Regn. No.: [REG_NO]")
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# ── 10. "Under the Supervision of" ───────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text Thesis 2"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run("Under the Supervision of")
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# ── 11. Guide name ────────────────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text 2 Doctor of Philosophy"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run("Dr. Rathnakar Achary")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(13)

# ── 12. Guide designation ─────────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text Thesis 2"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run("Associate Professor")
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# ── 13. Guide school ─────────────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text Thesis 2"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(36)
run = p.add_run("Alliance School of Advance Computing")
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# ── 14. University address ───────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text Thesis 2"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run("Alliance University")
run.font.name = "Times New Roman"
run.font.size = Pt(12)

p = doc.add_paragraph(style=sty("Text Thesis 2"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run("Central Campus, Chikkahadage Cross, Chandapura-Anekal Main Road, Bengaluru, Karnataka 562106")
run.font.name = "Times New Roman"
run.font.size = Pt(11)

# ── 15. Year ─────────────────────────────────────────────────────────────────
p = doc.add_paragraph(style=sty("Text Thesis 2"))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after  = Pt(0)
run = p.add_run("2026")
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(13)

# ── save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
